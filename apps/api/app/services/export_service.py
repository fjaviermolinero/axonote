"""
Servicio de exportación multi-modal para AxoNote.
Permite exportar contenido procesado en múltiples formatos académicos y profesionales.
"""

import asyncio
import json
import logging
import os
import tempfile
import zipfile
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union
from uuid import UUID, uuid4
import csv
import io

from sqlalchemy.orm import Session
from sqlalchemy import and_, or_

from app.core.config import get_settings
from app.models import (
    ClassSession, ExportSession, OCRResult, MicroMemo, MicroMemoCollection,
    TranscriptionResult, LLMAnalysisResult, ResearchResult, MedicalTerminology
)
from app.services.base import BaseService, ServiceConfigurationError
from app.services.minio_service import minio_service

logger = logging.getLogger(__name__)


class ConfiguracionExport:
    """Configuración para exportación de contenido."""
    
    def __init__(
        self,
        # Filtros de contenido
        incluir_transcripciones: bool = True,
        incluir_ocr: bool = True,
        incluir_micromemos: bool = True,
        incluir_research: bool = True,
        incluir_analytics: bool = False,
        
        # Filtros de calidad
        confianza_minima: float = 0.7,
        solo_validados: bool = False,
        
        # Filtros temporales
        fecha_inicio: Optional[datetime] = None,
        fecha_fin: Optional[datetime] = None,
        
        # Filtros especializados
        especialidades: List[str] = None,
        niveles_dificultad: List[str] = None,
        tipos_contenido: List[str] = None,
        
        # Opciones de formato
        incluir_metadatos: bool = True,
        incluir_imagenes: bool = True,
        incluir_audio: bool = False,
        comprimir_salida: bool = True,
        
        # Template y personalización
        template_personalizado: Optional[str] = None,
        logo_institucional: Optional[str] = None,
        header_personalizado: Optional[str] = None,
        footer_personalizado: Optional[str] = None,
        
        # Formato específico
        formato_referencias: str = "apa",  # apa, vancouver, harvard
        estilo_medico: str = "academico",  # clinico, academico, investigacion
        incluir_disclaimer: bool = True,
        confidencialidad: str = "medical"  # public, internal, medical, confidential
    ):
        self.incluir_transcripciones = incluir_transcripciones
        self.incluir_ocr = incluir_ocr
        self.incluir_micromemos = incluir_micromemos
        self.incluir_research = incluir_research
        self.incluir_analytics = incluir_analytics
        
        self.confianza_minima = confianza_minima
        self.solo_validados = solo_validados
        
        self.fecha_inicio = fecha_inicio
        self.fecha_fin = fecha_fin
        
        self.especialidades = especialidades or []
        self.niveles_dificultad = niveles_dificultad or []
        self.tipos_contenido = tipos_contenido or []
        
        self.incluir_metadatos = incluir_metadatos
        self.incluir_imagenes = incluir_imagenes
        self.incluir_audio = incluir_audio
        self.comprimir_salida = comprimir_salida
        
        self.template_personalizado = template_personalizado
        self.logo_institucional = logo_institucional
        self.header_personalizado = header_personalizado
        self.footer_personalizado = footer_personalizado
        
        self.formato_referencias = formato_referencias
        self.estilo_medico = estilo_medico
        self.incluir_disclaimer = incluir_disclaimer
        self.confidencialidad = confidencialidad


class ContenidoExport:
    """Contenedor para contenido agregado para export."""
    
    def __init__(self):
        self.transcripciones: List[Dict[str, Any]] = []
        self.ocr_results: List[Dict[str, Any]] = []
        self.micro_memos: List[Dict[str, Any]] = []
        self.research_results: List[Dict[str, Any]] = []
        self.analytics: Dict[str, Any] = {}
        self.metadatos: Dict[str, Any] = {}
    
    @property
    def total_elementos(self) -> int:
        """Número total de elementos de contenido."""
        return (
            len(self.transcripciones) +
            len(self.ocr_results) +
            len(self.micro_memos) +
            len(self.research_results)
        )
    
    def add_transcripcion(self, transcripcion: TranscriptionResult):
        """Añade transcripción al contenido."""
        self.transcripciones.append({
            "id": str(transcripcion.id),
            "texto": transcripcion.segments_with_speakers,
            "duracion": transcripcion.duration_seconds,
            "confianza": transcripcion.average_confidence,
            "idioma": transcripcion.language,
            "speakers": transcripcion.num_speakers,
            "created_at": transcripcion.created_at.isoformat()
        })
    
    def add_ocr_result(self, ocr_result: OCRResult):
        """Añade resultado OCR al contenido."""
        self.ocr_results.append({
            "id": str(ocr_result.id),
            "texto_original": ocr_result.extracted_text,
            "texto_corregido": ocr_result.corrected_text,
            "confianza": ocr_result.confidence_score,
            "tipo_contenido": ocr_result.content_type,
            "especialidad": ocr_result.medical_specialty,
            "terminos_medicos": ocr_result.medical_terms_detected,
            "created_at": ocr_result.created_at.isoformat()
        })
    
    def add_micro_memo(self, memo: MicroMemo):
        """Añade micro-memo al contenido."""
        self.micro_memos.append({
            "id": str(memo.id),
            "titulo": memo.title,
            "pregunta": memo.question,
            "respuesta": memo.answer,
            "explicacion": memo.explanation,
            "tipo": memo.memo_type,
            "dificultad": memo.difficulty_level,
            "confianza": memo.confidence_score,
            "tags": memo.tags,
            "fuente": memo.source_type,
            "created_at": memo.created_at.isoformat()
        })
    
    def add_research_result(self, research: ResearchResult):
        """Añade resultado de research al contenido."""
        self.research_results.append({
            "id": str(research.id),
            "query": research.research_query,
            "fuente": research.source_name,
            "contenido": research.content,
            "relevancia": research.relevance_score,
            "confianza": research.confidence_score,
            "created_at": research.created_at.isoformat()
        })


class ExportService(BaseService):
    """
    Servicio completo de exportación multi-modal.
    
    Exporta contenido procesado de AxoNote en múltiples formatos
    académicos y profesionales para uso educativo y clínico.
    """
    
    def __init__(self):
        super().__init__("ExportService")
        self.settings = get_settings()
        self.is_initialized = False
        
        # Configuración de paths
        self.export_base_path = Path(self.settings.EXPORT_STORAGE_PATH)
        self.templates_path = Path(self.settings.EXPORT_TEMPLATES_PATH)
        
        # Validación de configuración
        if not self.settings.EXPORT_ENABLED:
            raise ServiceConfigurationError("Export service is disabled")
    
    async def _setup(self):
        """Inicializa el servicio de export."""
        if self.is_initialized:
            return
        
        try:
            # Crear directorios necesarios
            self.export_base_path.mkdir(parents=True, exist_ok=True)
            
            # Validar templates disponibles
            await self._validate_templates()
            
            # Validar engines de export
            await self._validate_export_engines()
            
            self.is_initialized = True
            logger.info("Export service initialized successfully")
            
        except Exception as e:
            logger.error(f"Failed to initialize export service: {e}")
            raise ServiceConfigurationError(f"Export service setup failed: {e}")
    
    async def _validate_templates(self):
        """Valida que los templates necesarios estén disponibles."""
        required_templates = [
            "medical_academic.html",
            "clinical_report.html", 
            "study_cards.html",
            "anki_template.css",
            "export_styles.css"
        ]
        
        for template in required_templates:
            template_path = self.templates_path / template
            if not template_path.exists():
                logger.warning(f"Template not found: {template_path}")
    
    async def _validate_export_engines(self):
        """Valida que los engines de export estén disponibles."""
        validation_results = {}
        
        # Validar WeasyPrint para PDF
        try:
            import weasyprint
            validation_results["pdf"] = True
        except ImportError:
            validation_results["pdf"] = False
            logger.warning("WeasyPrint not available - PDF export disabled")
        
        # Validar python-docx para DOCX
        try:
            import docx
            validation_results["docx"] = True
        except ImportError:
            validation_results["docx"] = False
            logger.warning("python-docx not available - DOCX export disabled")
        
        # Validar genanki para Anki
        try:
            import genanki
            validation_results["anki"] = True
        except ImportError:
            validation_results["anki"] = False
            logger.warning("genanki not available - Anki export disabled")
        
        return validation_results
    
    async def health_check(self) -> Dict[str, Any]:
        """Verifica el estado de salud del servicio de export."""
        if not self.is_initialized:
            await self._setup()
        
        validation_results = await self._validate_export_engines()
        
        return {
            "service": "ExportService",
            "status": "healthy",
            "initialized": self.is_initialized,
            "export_engines": validation_results,
            "storage_path": str(self.export_base_path),
            "templates_path": str(self.templates_path),
            "supported_formats": [
                format_name for format_name, available in validation_results.items() if available
            ],
            "timestamp": datetime.utcnow().isoformat()
        }
    
    async def create_export_session(
        self,
        class_session_id: UUID,
        export_format: str,
        config: ConfiguracionExport,
        db: Session
    ) -> ExportSession:
        """
        Crea una nueva sesión de export.
        
        Args:
            class_session_id: ID de la sesión de clase
            export_format: Formato de exportación
            config: Configuración del export
            db: Sesión de base de datos
            
        Returns:
            ExportSession creada
        """
        if not self.is_initialized:
            await self._setup()
        
        try:
            # Verificar que la clase existe
            class_session = db.query(ClassSession).filter(
                ClassSession.id == class_session_id
            ).first()
            
            if not class_session:
                raise ValueError(f"Class session not found: {class_session_id}")
            
            # Crear sesión de export
            export_session = ExportSession(
                class_session_id=class_session_id,
                export_format=export_format,
                template_name=config.template_personalizado or "medical_academic",
                filters_applied=self._serialize_config(config),
                export_config={
                    "format": export_format,
                    "template": config.template_personalizado or "medical_academic",
                    "quality_threshold": config.confianza_minima,
                    "include_metadata": config.incluir_metadatos
                },
                export_title=f"Export {class_session.asignatura} - {class_session.tema}",
                export_description=f"Contenido exportado de la clase del {class_session.fecha}",
                language="ita",
                include_tts=config.incluir_audio
            )
            
            db.add(export_session)
            db.commit()
            db.refresh(export_session)
            
            logger.info(f"Created export session {export_session.id} for class {class_session_id}")
            return export_session
            
        except Exception as e:
            logger.error(f"Failed to create export session: {e}")
            db.rollback()
            raise
    
    async def gather_content_for_export(
        self,
        export_session: ExportSession,
        db: Session
    ) -> ContenidoExport:
        """
        Recopila todo el contenido para export según la configuración.
        
        Args:
            export_session: Sesión de export
            db: Sesión de base de datos
            
        Returns:
            ContenidoExport con todo el contenido filtrado
        """
        config = self._deserialize_config(export_session.filters_applied)
        contenido = ContenidoExport()
        
        # Metadatos de la clase
        class_session = db.query(ClassSession).filter(
            ClassSession.id == export_session.class_session_id
        ).first()
        
        contenido.metadatos = {
            "clase": {
                "id": str(class_session.id),
                "fecha": class_session.fecha.isoformat(),
                "asignatura": class_session.asignatura,
                "tema": class_session.tema,
                "profesor": class_session.profesor_text,
                "duracion": class_session.duracion_sec
            },
            "export": {
                "formato": export_session.export_format,
                "created_at": export_session.created_at.isoformat(),
                "filtros": export_session.filters_applied
            }
        }
        
        # Recopilar transcripciones
        if config.incluir_transcripciones:
            transcripciones = db.query(TranscriptionResult).filter(
                TranscriptionResult.class_session_id == export_session.class_session_id
            ).all()
            
            for trans in transcripciones:
                if trans.average_confidence >= config.confianza_minima:
                    contenido.add_transcripcion(trans)
        
        # Recopilar resultados OCR
        if config.incluir_ocr:
            ocr_results = db.query(OCRResult).filter(
                OCRResult.class_session_id == export_session.class_session_id
            ).all()
            
            for ocr in ocr_results:
                if ocr.confidence_score >= config.confianza_minima:
                    # Filtrar por especialidad si se especifica
                    if config.especialidades and ocr.medical_specialty not in config.especialidades:
                        continue
                    contenido.add_ocr_result(ocr)
        
        # Recopilar micro-memos
        if config.incluir_micromemos:
            memos = db.query(MicroMemo).filter(
                MicroMemo.class_session_id == export_session.class_session_id
            ).all()
            
            for memo in memos:
                if memo.confidence_score >= config.confianza_minima:
                    # Filtrar por dificultad si se especifica
                    if config.niveles_dificultad and memo.difficulty_level not in config.niveles_dificultad:
                        continue
                    contenido.add_micro_memo(memo)
        
        # Recopilar research results
        if config.incluir_research:
            research_results = db.query(ResearchResult).filter(
                ResearchResult.class_session_id == export_session.class_session_id
            ).all()
            
            for research in research_results:
                if research.confidence_score >= config.confianza_minima:
                    contenido.add_research_result(research)
        
        # Analytics si se solicita
        if config.incluir_analytics:
            contenido.analytics = await self._generate_analytics(
                export_session.class_session_id, db
            )
        
        logger.info(f"Gathered {contenido.total_elementos} elements for export {export_session.id}")
        return contenido
    
    async def export_to_pdf(
        self,
        export_session: ExportSession,
        contenido: ContenidoExport
    ) -> str:
        """
        Exporta contenido a PDF médico académico.
        
        Args:
            export_session: Sesión de export
            contenido: Contenido a exportar
            
        Returns:
            Path del archivo PDF generado
        """
        try:
            import weasyprint
            from jinja2 import Template
            
            # Cargar template
            template_path = self.templates_path / "medical_academic.html"
            with open(template_path, 'r', encoding='utf-8') as f:
                template_content = f.read()
            
            template = Template(template_content)
            
            # Generar HTML
            html_content = template.render(
                export_session=export_session,
                contenido=contenido,
                metadatos=contenido.metadatos,
                fecha_generacion=datetime.utcnow(),
                include_toc=True,
                include_metadata=True
            )
            
            # Configurar CSS
            css_path = self.templates_path / "export_styles.css"
            
            # Generar PDF
            output_path = self.export_base_path / f"{export_session.id}" / "export.pdf"
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            weasyprint.HTML(string=html_content, base_url=str(self.templates_path)).write_pdf(
                str(output_path),
                stylesheets=[str(css_path)] if css_path.exists() else None
            )
            
            logger.info(f"Generated PDF export: {output_path}")
            return str(output_path)
            
        except ImportError:
            raise ServiceConfigurationError("WeasyPrint not available for PDF export")
        except Exception as e:
            logger.error(f"Failed to generate PDF: {e}")
            raise
    
    async def export_to_docx(
        self,
        export_session: ExportSession,
        contenido: ContenidoExport
    ) -> str:
        """
        Exporta contenido a DOCX profesional.
        
        Args:
            export_session: Sesión de export
            contenido: Contenido a exportar
            
        Returns:
            Path del archivo DOCX generado
        """
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # Crear documento
            doc = Document()
            
            # Header del documento
            title = doc.add_heading(export_session.export_title or "Export Médico", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadatos
            doc.add_heading("Información de la Clase", level=1)
            meta_table = doc.add_table(rows=5, cols=2)
            meta_table.style = 'Table Grid'
            
            meta_data = [
                ("Asignatura", contenido.metadatos["clase"]["asignatura"]),
                ("Tema", contenido.metadatos["clase"]["tema"]),
                ("Fecha", contenido.metadatos["clase"]["fecha"]),
                ("Profesor", contenido.metadatos["clase"]["profesor"]),
                ("Formato Export", export_session.export_format.upper())
            ]
            
            for i, (key, value) in enumerate(meta_data):
                meta_table.cell(i, 0).text = key
                meta_table.cell(i, 1).text = str(value)
            
            # Contenido por secciones
            if contenido.transcripciones:
                doc.add_heading("Transcripciones", level=1)
                for trans in contenido.transcripciones:
                    doc.add_heading(f"Transcripción (Confianza: {trans['confianza']:.2f})", level=2)
                    doc.add_paragraph(trans["texto"])
            
            if contenido.ocr_results:
                doc.add_heading("Contenido OCR", level=1)
                for ocr in contenido.ocr_results:
                    doc.add_heading(f"OCR - {ocr['tipo_contenido']}", level=2)
                    doc.add_paragraph(ocr["texto_corregido"] or ocr["texto_original"])
            
            if contenido.micro_memos:
                doc.add_heading("Micro-Memos de Estudio", level=1)
                for memo in contenido.micro_memos:
                    doc.add_heading(memo["titulo"], level=2)
                    doc.add_paragraph(f"Pregunta: {memo['pregunta']}")
                    doc.add_paragraph(f"Respuesta: {memo['respuesta']}")
                    if memo["explicacion"]:
                        doc.add_paragraph(f"Explicación: {memo['explicacion']}")
                    doc.add_paragraph(f"Tipo: {memo['tipo']} | Dificultad: {memo['dificultad']}")
            
            if contenido.research_results:
                doc.add_heading("Referencias y Research", level=1)
                for research in contenido.research_results:
                    doc.add_heading(f"Fuente: {research['fuente']}", level=2)
                    doc.add_paragraph(research["contenido"])
            
            # Guardar archivo
            output_path = self.export_base_path / f"{export_session.id}" / "export.docx"
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            doc.save(str(output_path))
            
            logger.info(f"Generated DOCX export: {output_path}")
            return str(output_path)
            
        except ImportError:
            raise ServiceConfigurationError("python-docx not available for DOCX export")
        except Exception as e:
            logger.error(f"Failed to generate DOCX: {e}")
            raise
    
    async def export_to_json(
        self,
        export_session: ExportSession,
        contenido: ContenidoExport
    ) -> str:
        """
        Exporta contenido a JSON completo.
        
        Args:
            export_session: Sesión de export
            contenido: Contenido a exportar
            
        Returns:
            Path del archivo JSON generado
        """
        try:
            # Estructura JSON completa
            json_data = {
                "export_info": {
                    "id": str(export_session.id),
                    "formato": export_session.export_format,
                    "created_at": export_session.created_at.isoformat(),
                    "version": "1.0",
                    "generator": "AxoNote Export Service"
                },
                "metadata": contenido.metadatos,
                "content": {
                    "transcripciones": contenido.transcripciones,
                    "ocr_results": contenido.ocr_results,
                    "micro_memos": contenido.micro_memos,
                    "research_results": contenido.research_results
                },
                "analytics": contenido.analytics if contenido.analytics else {},
                "stats": {
                    "total_elementos": contenido.total_elementos,
                    "transcripciones_count": len(contenido.transcripciones),
                    "ocr_count": len(contenido.ocr_results),
                    "memos_count": len(contenido.micro_memos),
                    "research_count": len(contenido.research_results)
                }
            }
            
            # Guardar archivo
            output_path = self.export_base_path / f"{export_session.id}" / "export.json"
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(json_data, f, indent=2, ensure_ascii=False)
            
            logger.info(f"Generated JSON export: {output_path}")
            return str(output_path)
            
        except Exception as e:
            logger.error(f"Failed to generate JSON: {e}")
            raise
    
    async def export_to_csv(
        self,
        export_session: ExportSession,
        contenido: ContenidoExport
    ) -> str:
        """
        Exporta contenido a CSV para análisis.
        
        Args:
            export_session: Sesión de export
            contenido: Contenido a exportar
            
        Returns:
            Path del archivo CSV generado
        """
        try:
            output_path = self.export_base_path / f"{export_session.id}" / "export.csv"
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            with open(output_path, 'w', newline='', encoding='utf-8') as csvfile:
                writer = csv.writer(csvfile)
                
                # Headers
                writer.writerow([
                    "tipo_contenido", "id", "titulo", "contenido", "confianza",
                    "dificultad", "especialidad", "tags", "created_at"
                ])
                
                # Transcripciones
                for trans in contenido.transcripciones:
                    writer.writerow([
                        "transcripcion", trans["id"], "Transcripción",
                        trans["texto"][:500] + "..." if len(trans["texto"]) > 500 else trans["texto"],
                        trans["confianza"], "", "", "", trans["created_at"]
                    ])
                
                # OCR Results
                for ocr in contenido.ocr_results:
                    writer.writerow([
                        "ocr", ocr["id"], f"OCR - {ocr['tipo_contenido']}",
                        (ocr["texto_corregido"] or ocr["texto_original"])[:500],
                        ocr["confianza"], "", ocr["especialidad"], "", ocr["created_at"]
                    ])
                
                # Micro-memos
                for memo in contenido.micro_memos:
                    writer.writerow([
                        "micro_memo", memo["id"], memo["titulo"],
                        f"P: {memo['pregunta']} R: {memo['respuesta']}",
                        memo["confianza"], memo["dificultad"], "",
                        ";".join(memo["tags"]), memo["created_at"]
                    ])
                
                # Research
                for research in contenido.research_results:
                    writer.writerow([
                        "research", research["id"], f"Research - {research['fuente']}",
                        research["contenido"][:500],
                        research["confianza"], "", "", "", research["created_at"]
                    ])
            
            logger.info(f"Generated CSV export: {output_path}")
            return str(output_path)
            
        except Exception as e:
            logger.error(f"Failed to generate CSV: {e}")
            raise
    
    async def export_to_html(
        self,
        export_session: ExportSession,
        contenido: ContenidoExport
    ) -> str:
        """
        Exporta contenido a HTML interactivo.
        
        Args:
            export_session: Sesión de export
            contenido: Contenido a exportar
            
        Returns:
            Path del archivo HTML generado
        """
        try:
            from jinja2 import Template
            
            # Cargar template HTML interactivo
            template_path = self.templates_path / "interactive_report.html"
            
            # Template básico si no existe el archivo
            template_content = """
            <!DOCTYPE html>
            <html lang="it">
            <head>
                <meta charset="UTF-8">
                <meta name="viewport" content="width=device-width, initial-scale=1.0">
                <title>{{ export_session.export_title }}</title>
                <style>
                    body { font-family: Arial, sans-serif; margin: 20px; }
                    .header { background: #f4f4f4; padding: 20px; border-radius: 5px; }
                    .section { margin: 20px 0; }
                    .content-item { border: 1px solid #ddd; margin: 10px 0; padding: 15px; border-radius: 5px; }
                    .meta { color: #666; font-size: 0.9em; }
                    .search-box { width: 100%; padding: 10px; margin: 20px 0; }
                    .hidden { display: none; }
                </style>
            </head>
            <body>
                <div class="header">
                    <h1>{{ export_session.export_title }}</h1>
                    <p>Generado el {{ fecha_generacion.strftime('%d/%m/%Y %H:%M') }}</p>
                    <p>Total elementos: {{ contenido.total_elementos }}</p>
                </div>
                
                <input type="text" class="search-box" placeholder="Buscar en el contenido..." onkeyup="filtrarContenido(this.value)">
                
                {% if contenido.transcripciones %}
                <div class="section">
                    <h2>Transcripciones ({{ contenido.transcripciones|length }})</h2>
                    {% for trans in contenido.transcripciones %}
                    <div class="content-item searchable" data-content="{{ trans.texto|lower }}">
                        <h3>Transcripción</h3>
                        <p>{{ trans.texto }}</p>
                        <div class="meta">Confianza: {{ trans.confianza }}, Duración: {{ trans.duracion }}s</div>
                    </div>
                    {% endfor %}
                </div>
                {% endif %}
                
                {% if contenido.micro_memos %}
                <div class="section">
                    <h2>Micro-Memos ({{ contenido.micro_memos|length }})</h2>
                    {% for memo in contenido.micro_memos %}
                    <div class="content-item searchable" data-content="{{ (memo.titulo + ' ' + memo.pregunta + ' ' + memo.respuesta)|lower }}">
                        <h3>{{ memo.titulo }}</h3>
                        <p><strong>Pregunta:</strong> {{ memo.pregunta }}</p>
                        <p><strong>Respuesta:</strong> {{ memo.respuesta }}</p>
                        {% if memo.explicacion %}<p><strong>Explicación:</strong> {{ memo.explicacion }}</p>{% endif %}
                        <div class="meta">Tipo: {{ memo.tipo }}, Dificultad: {{ memo.dificultad }}</div>
                    </div>
                    {% endfor %}
                </div>
                {% endif %}
                
                <script>
                function filtrarContenido(query) {
                    const items = document.querySelectorAll('.searchable');
                    query = query.toLowerCase();
                    
                    items.forEach(item => {
                        const content = item.getAttribute('data-content');
                        if (content.includes(query)) {
                            item.classList.remove('hidden');
                        } else {
                            item.classList.add('hidden');
                        }
                    });
                }
                </script>
            </body>
            </html>
            """
            
            if template_path.exists():
                with open(template_path, 'r', encoding='utf-8') as f:
                    template_content = f.read()
            
            template = Template(template_content)
            
            # Generar HTML
            html_content = template.render(
                export_session=export_session,
                contenido=contenido,
                fecha_generacion=datetime.utcnow()
            )
            
            # Guardar archivo
            output_path = self.export_base_path / f"{export_session.id}" / "export.html"
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            logger.info(f"Generated HTML export: {output_path}")
            return str(output_path)
            
        except Exception as e:
            logger.error(f"Failed to generate HTML: {e}")
            raise
    
    async def export_to_anki(
        self,
        export_session: ExportSession,
        contenido: ContenidoExport
    ) -> str:
        """
        Exporta micro-memos a package Anki (.apkg).
        
        Args:
            export_session: Sesión de export
            contenido: Contenido a exportar
            
        Returns:
            Path del archivo .apkg generado
        """
        try:
            import genanki
            import random
            
            # Crear modelo de nota Anki
            modelo = genanki.Model(
                random.randrange(1 << 30, 1 << 31),
                'Modelo Médico AxoNote',
                fields=[
                    {'name': 'Pregunta'},
                    {'name': 'Respuesta'},
                    {'name': 'Explicacion'},
                    {'name': 'Tipo'},
                    {'name': 'Dificultad'},
                    {'name': 'Tags'}
                ],
                templates=[
                    {
                        'name': 'Tarjeta Médica',
                        'qfmt': '''
                        <div class="pregunta">{{Pregunta}}</div>
                        <div class="meta">Tipo: {{Tipo}} | Dificultad: {{Dificultad}}</div>
                        ''',
                        'afmt': '''
                        <div class="pregunta">{{Pregunta}}</div>
                        <hr>
                        <div class="respuesta">{{Respuesta}}</div>
                        {{#Explicacion}}<div class="explicacion"><strong>Explicación:</strong> {{Explicacion}}</div>{{/Explicacion}}
                        <div class="meta">{{Tags}}</div>
                        '''
                    }
                ],
                css='''
                .pregunta { font-size: 18px; font-weight: bold; margin-bottom: 10px; }
                .respuesta { font-size: 16px; color: #2c5aa0; }
                .explicacion { font-size: 14px; margin-top: 10px; background: #f0f8ff; padding: 10px; }
                .meta { font-size: 12px; color: #666; margin-top: 10px; }
                '''
            )
            
            # Crear deck
            deck = genanki.Deck(
                random.randrange(1 << 30, 1 << 31),
                f"AxoNote - {contenido.metadatos['clase']['asignatura']}"
            )
            
            # Añadir micro-memos como notas
            for memo in contenido.micro_memos:
                nota = genanki.Note(
                    model=modelo,
                    fields=[
                        memo["pregunta"],
                        memo["respuesta"],
                        memo["explicacion"] or "",
                        memo["tipo"],
                        memo["dificultad"],
                        ", ".join(memo["tags"])
                    ],
                    tags=memo["tags"]
                )
                deck.add_note(nota)
            
            # Crear package
            package = genanki.Package(deck)
            
            # Guardar archivo
            output_path = self.export_base_path / f"{export_session.id}" / "export.apkg"
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            package.write_to_file(str(output_path))
            
            logger.info(f"Generated Anki package: {output_path}")
            return str(output_path)
            
        except ImportError:
            raise ServiceConfigurationError("genanki not available for Anki export")
        except Exception as e:
            logger.error(f"Failed to generate Anki package: {e}")
            raise
    
    def _serialize_config(self, config: ConfiguracionExport) -> Dict[str, Any]:
        """Serializa configuración para almacenamiento."""
        return {
            "incluir_transcripciones": config.incluir_transcripciones,
            "incluir_ocr": config.incluir_ocr,
            "incluir_micromemos": config.incluir_micromemos,
            "incluir_research": config.incluir_research,
            "incluir_analytics": config.incluir_analytics,
            "confianza_minima": config.confianza_minima,
            "solo_validados": config.solo_validados,
            "especialidades": config.especialidades,
            "niveles_dificultad": config.niveles_dificultad,
            "incluir_metadatos": config.incluir_metadatos,
            "incluir_imagenes": config.incluir_imagenes,
            "incluir_audio": config.incluir_audio,
            "comprimir_salida": config.comprimir_salida
        }
    
    def _deserialize_config(self, config_dict: Dict[str, Any]) -> ConfiguracionExport:
        """Deserializa configuración desde almacenamiento."""
        return ConfiguracionExport(**config_dict)
    
    async def _generate_analytics(
        self,
        class_session_id: UUID,
        db: Session
    ) -> Dict[str, Any]:
        """Genera analytics del contenido de la clase."""
        # Implementación básica de analytics
        analytics = {
            "contenido_stats": {},
            "calidad_promedio": 0.0,
            "distribucion_tipos": {},
            "timeline": []
        }
        
        # TODO: Implementar analytics completas
        
        return analytics


# Instancia singleton del servicio
export_service = ExportService()
